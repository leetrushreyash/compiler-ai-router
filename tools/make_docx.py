from pathlib import Path
from docx import Document

ROOT = Path(__file__).resolve().parents[1]
docs_dir = ROOT / "code_smell_compiler" / "docs"
out_path = docs_dir / "Combined_Project_Docs.docx"

files = [
    "problem_statement.txt",
    "novelty.txt",
    "architecture.txt",
    "prerequisites.txt",
    "tools.txt",
]

doc = Document()
doc.add_heading('ML-Driven Code Smell Detection Compiler - Combined Documents', level=1)

for fname in files:
    p = docs_dir / fname
    if not p.exists():
        continue
    title = p.stem.replace('_', ' ').title()
    doc.add_heading(title, level=2)
    text = p.read_text(encoding='utf-8')
    # split into paragraphs by blank lines
    paras = [para.strip() for para in text.split('\n\n') if para.strip()]
    for para in paras:
        # preserve line breaks within paragraph
        lines = para.splitlines()
        if len(lines) == 1:
            doc.add_paragraph(lines[0])
        else:
            p_obj = doc.add_paragraph(lines[0])
            for ln in lines[1:]:
                p_obj.add_run('\n' + ln)

doc.save(out_path)
print(f"Saved combined docx to: {out_path}")
