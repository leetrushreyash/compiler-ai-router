from docx import Document
from pathlib import Path

srs_txt = Path(r"c:\Users\SHREYASH SHARMA\compiler\code_smell_compiler\docs\SRS.txt")
if not srs_txt.exists():
    raise SystemExit(f"SRS source not found: {srs_txt}")

doc = Document()
with srs_txt.open('r', encoding='utf-8') as f:
    lines = f.readlines()

for line in lines:
    stripped = line.strip()
    if stripped == '':
        doc.add_paragraph('')
        continue
    # simple heuristic: headings are lines that end with ':' or are numbered
    if stripped.startswith('1.') or stripped[:2].isdigit():
        doc.add_heading(stripped, level=2)
    else:
        doc.add_paragraph(stripped)

out = Path(r"c:\Users\SHREYASH SHARMA\compiler\code_smell_compiler\docs\SRS.docx")
doc.save(out)
print(f"Saved SRS doc to: {out}")
